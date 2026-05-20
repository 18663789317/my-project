# -*- coding: utf-8 -*-
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "CH场外期权系统_产品文档与交易台API对接需求_V3.6.docx"


def set_default_font(doc: Document, east_asia_font: str = "微软雅黑", latin_font: str = "Calibri") -> None:
    style = doc.styles["Normal"]
    style.font.name = latin_font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)
    style.font.size = Pt(10.5)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for t in items:
        doc.add_paragraph(str(t), style="List Bullet")


def add_numbered(doc: Document, items: Sequence[str]) -> None:
    for t in items:
        doc.add_paragraph(str(t), style="List Number")


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = "Table Grid"
    head = tb.rows[0].cells
    for i, h in enumerate(headers):
        head[i].text = str(h)
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)


def build_document() -> Document:
    doc = Document()
    set_default_font(doc)

    doc.add_heading("CH场外期权结构风险管理监控系统 V3.6", level=0)
    p = doc.add_paragraph()
    p.add_run("产品文档 + 交易台API对接需求（详细版）\n").bold = True
    p.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
    p.add_run("适用范围：当前 app.py 实现版本（单文件 Streamlit + SQLite 架构）。")

    doc.add_heading("1. 文档目的", level=1)
    doc.add_paragraph(
        "本文件用于统一业务、产品、开发、测试、实施和交易台对接方的理解，覆盖系统能力边界、数据口径、自动化流程、"
        "以及与真实交易台打通所需的接口规范。"
    )

    doc.add_heading("2. 产品定位与价值", level=1)
    doc.add_paragraph(
        "系统定位为“结构建仓、价格维护、头寸平仓、风险监控、汇报导出”的一体化工作台，目标是把分散在交易、风控、运营端的"
        "链路统一到一个可追溯账本中。"
    )
    add_bullets(
        doc,
        [
            "统一口径：所有页面基于同一计算引擎（compute_ledgers）回算。",
            "统一数据：结构、价格、平仓、现货、风险暴露均入库可追溯。",
            "统一监控：日度监控、敞口区间、平仓收益、图片汇报联动输出。",
            "自动化增强：自动补价、自动补贴/雪球转换同步，降低手工维护成本。",
        ],
    )

    doc.add_heading("3. 系统架构概览", level=1)
    add_bullets(
        doc,
        [
            "前端框架：Streamlit",
            "数据存储：SQLite（otc_gui.db）",
            "计算引擎：Pandas + 策略状态机",
            "外部行情：AKShare（日线收盘价）",
            "图像输出：Matplotlib",
        ],
    )

    doc.add_heading("4. 菜单模块说明", level=1)
    add_table(
        doc,
        headers=["菜单", "目标", "关键能力"],
        rows=[
            ["生成策略组", "维护策略组主数据与迁移", "创建/更新、导入导出、冲突检测、编号重命名"],
            ["结构录入", "维护结构参数与生命周期", "新增/编辑、终止、报价图预览、删除治理"],
            ["价格录入", "维护收盘价完整性", "批量录入、单日录入、可编辑表、AK自动导入、整点自动补空值"],
            ["现货头寸仓库管理", "维护现货库存流水", "买入增仓、纯现货平仓、汇总与回溯"],
            ["期权头寸仓库管理", "维护结构在库与平仓", "快速平仓、对称平仓、现货对冲、外部平仓、批次撤回"],
            ["监控计算", "日度风险与收益监控", "结构/组/敞口/平仓监控、图片导出、K线展示"],
        ],
    )

    doc.add_heading("5. 核心数据模型", level=1)
    add_table(
        doc,
        headers=["表名", "用途", "主键/约束", "关键字段"],
        rows=[
            ["strategy_group", "策略组主数据", "group_id(PK)", "group_name, underlying"],
            ["structure", "结构定义", "structure_id(PK)", "group_id, kind, strategy_code, start/end, params_json"],
            ["price", "价格序列", "PRIMARY(dt, underlying)", "settle, source, is_locked, updated_at"],
            ["close_trade2", "平仓明细", "close_id(PK)", "dt, structure_id, side, qty, open/close_price, pnl, category"],
            ["snowball_conversion", "雪球折价转换", "UNIQUE(structure_id, trigger_date)", "conversion_qty, conversion_price, source_status"],
            ["spot_position_lot", "现货买入批次", "lot_id(PK)", "spot_name, buy_dt, qty, buy_price"],
            ["spot_hedge_match_log", "现货对冲流水", "match_id(PK)", "spot_lot_id, structure_id, matched_qty, total_pnl"],
            ["close_revert_log", "平仓撤回日志", "log_id(PK)", "close_id, payload_json, reverted_at"],
            ["app_kv", "系统KV配置", "k(PK)", "v, updated_at"],
        ],
    )

    doc.add_heading("6. 关键计算口径", level=1)
    add_bullets(
        doc,
        [
            "结构日度：按策略状态机逐交易日计算生成量、状态、当日/累计盈亏。",
            "策略组日度：聚合净持仓、均价、浮动盈亏、平仓盈亏、补贴盈亏、总盈亏。",
            "风险边界：输出结构级与组级的剩余最小/最大、敞口下界/上界。",
            "终止处理：手动终结、熔断终止、雪球转换会收敛剩余交易日与剩余规模。",
        ],
    )

    doc.add_heading("7. 已实现自动化（含新增）", level=1)
    add_bullets(
        doc,
        [
            "自动同步熔断补贴平仓记录（幂等）。",
            "自动同步雪球折价转期货记录（幂等）。",
            "价格页整点自动补空值（交易日15:00后，每整点每策略组最多一次）。",
        ],
    )
    doc.add_paragraph("价格页整点自动补空值规则：")
    add_numbered(
        doc,
        [
            "时区按中国时间（Asia/Shanghai）判定。",
            "仅交易日且本地时间 >= 15:00 才尝试。",
            "频率：同策略组同小时只执行一次（例如15点一次、16点一次）。",
            "范围：仅当前策略组相关品种（策略组默认品种 + 该组结构涉及品种）。",
            "策略：只补空值（缺失记录或空settle），不覆盖已有有效值，保护手工锁定。",
            "提醒：仅在“实际写入成功 > 0”时显示顶部轻提示，无更新不提示。",
        ],
    )

    doc.add_heading("8. 与交易台打通的API需求（重点）", level=1)
    doc.add_paragraph("建议按“P0必需 + P1增强”分层推进。")

    doc.add_heading("8.1 P0 必需API清单", level=2)
    add_table(
        doc,
        headers=["API", "目的", "最少字段", "更新频率"],
        rows=[
            [
                "Auth API",
                "鉴权与token获取",
                "client_id, token, expire_at",
                "按需",
            ],
            [
                "交易日历 API",
                "统一交易日判断",
                "trade_date, is_trading_day, exchange",
                "每日/按需",
            ],
            [
                "结构成交单 API",
                "自动录入新建结构",
                "external_trade_id, book/group, strategy_code, kind, underlying, start/end, entry/strike/barrier, base_qty_per_day, risk_party, status, version, updated_at",
                "实时+增量",
            ],
            [
                "结构变更/终止 API",
                "自动更新结构状态",
                "external_trade_id, event_type(AMEND/TERMINATE/CANCEL), effective_date, version, updated_at",
                "实时+增量",
            ],
            [
                "平仓明细 API",
                "自动落地平仓流水",
                "close_id, external_trade_id, close_date, side, qty, open_price, close_price, pnl, close_category, is_external, batch_id, updated_at",
                "实时+增量",
            ],
            [
                "官方收盘价 API",
                "每日价格入库",
                "trade_date, underlying, close_price, source, updated_at",
                "每日EOD",
            ],
            [
                "头寸快照 API",
                "对账每日持仓",
                "as_of_date, external_trade_id, long_qty, short_qty, net_qty, avg_price, status",
                "每日",
            ],
            [
                "风险快照 API",
                "对账与监控风险暴露",
                "as_of_date, external_trade_id, delta, gamma, vega, theta, notional, exposure_min/max, limit_usage",
                "每日",
            ],
            [
                "标的主数据 API",
                "统一合约编码与乘数",
                "underlying, exchange, contract_multiplier, tick_size, active_flag",
                "每日/变更时",
            ],
            [
                "账户/簿记映射 API",
                "映射到策略组",
                "book_id, book_name, group_id, owner",
                "每日/变更时",
            ],
        ],
    )

    doc.add_heading("8.2 P1 增强API清单", level=2)
    add_bullets(
        doc,
        [
            "事件序列 API（since_seq）：保证增量不丢不重。",
            "全量快照 API：每日收盘后做全量对账修复。",
            "数据更正 API：处理撤单、改价、回补、反冲。",
            "回执 API：我方落库成功后回执交易台，形成闭环。",
        ],
    )

    doc.add_heading("8.3 对接契约（必须一次谈清）", level=2)
    add_numbered(
        doc,
        [
            "业务主键：external_trade_id、close_id必须全局唯一且不可复用。",
            "幂等规则：同业务键重复推送不得重复入账。",
            "版本控制：所有可变对象带version，只接受新版本覆盖。",
            "时间口径：统一Asia/Shanghai，明确到秒/毫秒。",
            "增量机制：支持updated_since或since_seq。",
            "更正机制：必须有event_type/is_correction字段。",
            "价格口径：明确close_price，不混用结算价/即时价。",
            "SLA：明确最晚到达时间（收盘价、头寸快照、风险快照）。",
            "错误码：统一code/message/retriable规范。",
            "限流重试：QPS、429回退、重试窗口和重放周期。",
        ],
    )

    doc.add_heading("8.4 与本系统表映射", level=2)
    add_table(
        doc,
        headers=["交易台数据", "本系统落地表", "说明"],
        rows=[
            ["结构成交单/变更", "structure", "主数据同步，按external_trade_id映射structure_id（建议保留外部ID字段）"],
            ["官方收盘价", "price", "交易日口径，只写收盘价"],
            ["平仓明细", "close_trade2", "逐笔明细入账，支持外部平仓标识"],
            ["雪球转换事件", "snowball_conversion", "若交易台提供事件则直接写入；否则由内部引擎同步"],
            ["账户与簿记映射", "strategy_group", "book/group映射到业务策略组"],
            ["头寸/风险快照", "监控页对账数据集", "可入独立快照表用于核对与审计"],
        ],
    )

    doc.add_heading("9. API字段字典建议（请求交易台时可直接引用）", level=1)
    doc.add_paragraph("以下字段建议统一命名和类型，减少后续联调歧义。")
    add_table(
        doc,
        headers=["字段名", "类型", "是否必填", "说明"],
        rows=[
            ["external_trade_id", "string", "是", "交易台结构唯一ID，不可复用"],
            ["close_id", "string", "是", "平仓明细唯一ID，不可复用"],
            ["version", "int", "是", "对象版本号，单调递增"],
            ["updated_at", "datetime", "是", "最后更新时间（中国时区）"],
            ["event_type", "string", "是", "NEW/AMEND/TERMINATE/CANCEL/CORRECTION"],
            ["underlying", "string", "是", "统一标的编码（如I2605）"],
            ["trade_date", "date", "是", "交易日（yyyy-mm-dd）"],
            ["close_price", "number(18,6)", "是", "收盘价，禁止结算价替代"],
            ["qty", "number(18,4)", "是", "数量，统一吨/手口径并注明"],
            ["pnl", "number(18,6)", "否", "平仓盈亏或事件盈亏"],
            ["is_correction", "bool", "否", "是否更正消息"],
            ["source_system", "string", "否", "来源系统标识（便于审计）"],
        ],
    )

    doc.add_heading("10. 联调与验收建议", level=1)
    add_numbered(
        doc,
        [
            "先打通P0中的结构、价格、平仓3条主链路。",
            "以1个策略组做7个交易日样本，核对结构状态、头寸、盈亏、敞口。",
            "验证幂等：重复推送同一批数据不应产生重复记录。",
            "验证更正：撤单/改单后目标记录应按version正确覆盖。",
            "验证时点：15:00后自动补价逻辑与交易台收盘价到达时点一致。",
            "验证监控：监控页面指标与交易台日报可对账。",
        ],
    )

    doc.add_heading("11. 部署与运维建议", level=1)
    add_bullets(
        doc,
        [
            "数据库与导出文件目录每日备份。",
            "关键作业（自动补价、自动同步）记录执行日志与结果计数。",
            "接口失败按可重试类型退避重试，不可重试进入告警清单。",
            "每周做一次交易台与本系统的全量对账。",
        ],
    )

    doc.add_heading("12. 附录：给交易台的需求清单（可直接发）", level=1)
    doc.add_paragraph("请交易台提供以下内容：")
    add_numbered(
        doc,
        [
            "API文档（OpenAPI/Swagger）+ 沙箱地址 + 联调账号。",
            "字段字典（类型、单位、枚举值、是否可空、默认值）。",
            "增量同步机制（updated_since 或 since_seq）与历史回放窗口。",
            "错误码与限流策略。",
            "交易日历与收盘价到达SLA。",
            "示例数据包（至少覆盖：新建、变更、终止、平仓、更正）。",
        ],
    )

    doc.add_paragraph("\n—— 文档结束 ——")
    return doc


def main() -> None:
    doc = build_document()
    doc.save(str(OUT))
    print(str(OUT))


if __name__ == "__main__":
    main()

