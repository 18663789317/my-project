from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "场外期权结构风险管理监控系统_用户使用手册.docx"
ALT_OUT_PATH = ROOT / "OTC_User_Manual_CN.docx"


def set_east_asia_font(style, font_name: str) -> None:
    style.font.name = font_name
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_run_font(run, font_name: str = "Microsoft YaHei", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 - level if level <= 3 else 11, bold=True)


def add_paragraph(doc: Document, text: str = "", *, style: str | None = None, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.25


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.25


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                set_run_font(r, bold=True)
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for p in cells[idx].paragraphs:
                for r in p.runs:
                    set_run_font(r)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph("")


def build_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.4)
    sec.right_margin = Cm(2.2)

    styles = doc.styles
    set_east_asia_font(styles["Normal"], "Microsoft YaHei")
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Number"]:
        if style_name in styles:
            set_east_asia_font(styles[style_name], "Microsoft YaHei")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("场外期权结构风险管理监控系统\n用户使用手册")
    set_run_font(r, size=22, bold=True)
    r.font.color.rgb = RGBColor(15, 43, 91)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("中文版 Word 文档\n适用范围：当前单文件版系统（含 TRS 监控、仓库管理、监控计算）")
    set_run_font(r, size=12)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = info.add_run(f"生成日期：{date.today().strftime('%Y-%m-%d')}")
    set_run_font(r, size=11)

    add_paragraph(doc, "")
    add_paragraph(doc, "本手册面向业务使用人员、交易员、套保人员、风控人员和运营人员。目标是让用户不看代码，也能理解系统的页面功能、推荐操作顺序、关键按钮用途和常见问题处理方式。")

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    footer = doc.sections[-1].footer.paragraphs[0]
    add_page_number(footer)
    return doc


def write_manual(doc: Document) -> None:
    add_heading(doc, "1. 文档说明", level=1)
    add_bullets(
        doc,
        [
            "手册定位：用于指导业务用户完成策略组建立、结构录入、价格维护、现货与期权仓库管理、平仓回溯、监控计算和日报图片输出。",
            "阅读方式：建议先阅读“建议操作顺序”，再按业务场景跳转到对应功能章节。",
            "日期口径：系统中大多数日期控件采用“YYYY/MM/DD”输入格式；监控口径受“监控日期”统一控制。",
            "编号口径：结构编号（structure_id）在当前系统里是全局唯一，不仅在本策略组内唯一。",
        ],
    )

    add_heading(doc, "2. 软件概览", level=1)
    add_paragraph(doc, "系统是一套面向场外期权、TRS、现货头寸和风险监控的业务工作台，覆盖从主数据建立、日常录入，到仓库管理、平仓、监控、汇报输出的完整流程。")
    add_table(
        doc,
        ["功能菜单", "主要用途", "典型使用者"],
        [
            ["生成策略组", "建立策略组主数据；导入/导出 JSON；处理编号冲突", "运营、管理员"],
            ["结构录入", "录入结构；下载结构报价图；维护当前组内结构", "交易、运营"],
            ["价格录入", "录入或批量维护收盘价；可选接入 AKShare", "运营、风控"],
            ["现货头寸仓库管理", "维护现货买入、现货平仓、明细流水和汇总清零", "现货交易、运营"],
            ["期权头寸仓库管理", "查看在库结构头寸、TRS 管理、各类平仓及平仓回溯", "交易、套保、运营"],
            ["监控计算", "执行监控口径计算，查看结构/风险/平仓结果，导出汇报图片", "风控、管理层、运营"],
        ],
    )

    add_heading(doc, "3. 上手前准备", level=1)
    add_numbers(
        doc,
        [
            "确认系统已打开到浏览器页面，左侧可以看到 6 个功能菜单：生成策略组、结构录入、价格录入、现货头寸仓库管理、期权头寸仓库管理、监控计算。",
            "先创建策略组，再录入结构；如果没有策略组或结构，后续多个页面会直接提示“请先创建策略组”或“请先创建策略组和结构”。",
            "若需要使用“AK 自动导入”或“快速更新今收”，需提前安装 AKShare；未安装时，手工价格录入功能仍可正常使用。",
            "若涉及现货对冲，请先录入现货头寸，再进行现货对冲平仓；否则相关页面不会出现可配对的现货。",
        ],
    )
    add_paragraph(doc, "关键业务概念：")
    add_table(
        doc,
        ["概念", "说明"],
        [
            ["策略组", "系统的业务主容器。结构、现货头寸、平仓记录、监控结果都基于策略组组织。"],
            ["结构", "期权或 TRS 业务单元。每条结构有结构编号、风险子、品种、方向、日期区间和策略参数。"],
            ["风险子", "结构所属的风险承担主体，如海证资本。"],
            ["品种", "标的合约代码或品种代码，例如 I2605、RB2605。"],
            ["仓库统计截至日期", "期权头寸仓库管理页面的仓库时间截面；会影响仓库头寸与平仓默认日期。"],
            ["监控日期", "监控计算页的统一回算日期；下方所有表和图片都以该日期为准。"],
            ["手工锁定价格", "价格页中可锁定某天价格，AK 自动写入时默认不会覆盖锁定记录。"],
        ],
    )

    add_heading(doc, "4. 建议操作顺序", level=1)
    add_numbers(
        doc,
        [
            "在“生成策略组”创建策略组主数据，确定策略组编号、名称、默认品种。",
            "在“结构录入”录入结构；如需报价图片，可在保存前或保存后下载结构报价 PNG。",
            "在“价格录入”维护历史价格和最新收盘价，确保监控和仓库计算有完整价格基础。",
            "如业务涉及现货，在“现货头寸仓库管理”录入现货买入或进行纯现货平仓。",
            "在“期权头寸仓库管理”查看在库头寸、执行结构平仓、TRS 新增/换月、多空对称平仓、现货对冲平仓。",
            "最后在“监控计算”统一查看风险结果，下载汇报图片，并根据需要导出 CSV。"],
    )

    add_heading(doc, "5. 生成策略组", level=1)
    add_paragraph(doc, "本页面用于建立和维护策略组主数据，也是数据导入导出的入口。")
    add_heading(doc, "5.1 创建或更新策略组", level=2)
    add_numbers(
        doc,
        [
            "在左侧菜单进入“生成策略组”。",
            "在左侧表单依次填写“策略组编号”“策略组名称”“默认品种”。",
            "点击“创建/更新策略组”。如果编号已存在，系统会按该编号直接更新名称和默认品种。",
            "保存成功后，可在右侧表格中看到该策略组。"],
    )
    add_heading(doc, "5.2 导出策略组", level=2)
    add_bullets(
        doc,
        [
            "在“策略组导入/导出”区域，使用“选择要导出的策略组（可多选）”选择一个或多个策略组。",
            "点击“导出选中策略组”，系统会导出 JSON 文件。",
            "导出包包含：策略组、现货头寸、结构、平仓汇总、平仓明细、现货对冲流水、撤回日志和价格。",
            "适用场景：迁移、备份、测试环境导入、跨电脑交接。"],
    )
    add_heading(doc, "5.3 导入策略组", level=2)
    add_numbers(
        doc,
        [
            "点击“导入策略组文件（JSON）”，选择导出的 JSON 文件。",
            "系统会先执行体检和冲突检查，包括：策略组冲突、结构 ID 冲突、平仓记录 ID 冲突、现货头寸 ID 冲突、现货对冲流水 ID 冲突。",
            "若策略组编号冲突，可选择两种方式：一是覆盖导入；二是在弹窗中重命名策略组编号。",
            "若是非策略组主键冲突，系统会自动重命名并给出映射摘要。",
            "确认“覆盖导入风险”或“价格全量覆盖风险”后，点击“确认导入策略组”。"],
    )
    add_paragraph(doc, "风险提示：覆盖导入会先删除旧数据，再导入新数据，属于不可自动回退操作。导入前建议先导出备份。")
    add_heading(doc, "5.4 策略组表格维护", level=2)
    add_bullets(
        doc,
        [
            "右侧表格支持按“策略组编号 / 策略组名称 / 默认品种”筛选。",
            "点击“保存策略组表格修改”可批量保存右侧表格编辑结果。",
            "勾选“删除”后点击“删除选中策略组”，系统会级联删除该策略组的关联数据。"],
    )

    add_heading(doc, "6. 结构录入", level=1)
    add_paragraph(doc, "本页面用于录入和维护期权结构，也可在保存时下载结构报价图片。右侧表格用于维护当前组内结构。")
    add_heading(doc, "6.1 左侧录入区的标准步骤", level=2)
    add_numbers(
        doc,
        [
            "先在顶部“选择策略组”中选定当前要维护的策略组。",
            "填写“结构编号”。系统会自动给出建议编号，但用户也可以手工修改；若该编号已被其他策略组占用，系统会阻止保存。",
            "选择“风险子”，填写“标的名称”，并在“标的代码”处选择实际合约代码。",
            "选择“期权结构”，例如普通累计、无敲出累计、浮动熔断累计、安全气囊、雪球、TRS。",
            "填写“报价日期”“方向”“开始日期”“结束日期”和该策略对应的参数。",
            "如需要对外沟通，可下载“结构报价图片 PNG”。",
            "点击“保存结构”；若系统检测到该结构在已录价格路径上会立即触发终止，会先弹出二次确认。"],
    )
    add_heading(doc, "6.2 特殊结构说明", level=2)
    add_bullets(
        doc,
        [
            "TRS：系统会自动给出默认开始日和结束日。TRS 录入时以一次性头寸数量为核心，后续主要在期权头寸仓库管理和监控计算中的“TRS监控”查看。",
            "雪球：会出现额外参数区，包括期限单位、期限期数、敲出观察频率、票息、敲入/敲出参数、折价转期货等。",
            "安全气囊：方向显示为看涨/看跌，系统在监控页会按参与率单独展示。",
            "同一条结构保存后，系统会参与后续价格回算、仓库生成、平仓和监控。"],
    )
    add_heading(doc, "6.3 右侧“当前组内结构”维护", level=2)
    add_bullets(
        doc,
        [
            "右侧表格默认展示当前组的存续结构。",
            "可使用“结构复筛选”对结构编号、品种、方向、风险子等进行快速过滤。",
            "点击“保存表格修改”可批量保存结构元数据修改。",
            "点击“删除选中结构”会删除结构及其关联的平仓记录、撤回日志和相关对冲记录。",
            "在“查看已终止结构（手动终结/熔断终止）”中，可对已终止结构执行“确认彻底删除”。",
        ],
    )
    add_paragraph(doc, "重要提示：结构删掉后，该结构会从活动页面消失；请确认不再需要保留历史业务关系后再执行删除。")

    add_heading(doc, "7. 价格录入", level=1)
    add_paragraph(doc, "价格页用于维护收盘价，是仓库、监控和日报计算的基础。没有价格，系统很多结果无法正确生成。")
    add_heading(doc, "7.1 按结构区间批量录入", level=2)
    add_numbers(
        doc,
        [
            "选择策略组后，在“按选定结构区间批量录入（横向日期）”区域选择结构。",
            "系统会按该结构的开始日到结束日展开交易日横向表格。",
            "在对应日期单元格中录入收盘价。",
            "点击“保存结构区间价格”。系统仅保存交易日价格，非交易日会被跳过。"],
    )
    add_heading(doc, "7.2 单日快速录入/修改", level=2)
    add_bullets(
        doc,
        [
            "在“单日快速录入/修改”区域，选择品种、日期和收盘价。",
            "点击“保存单日价格”即可写入或覆盖该天该品种价格。",
            "若日期不是交易日，系统会提示“该日期非交易日，未保存”。"],
    )
    add_heading(doc, "7.3 价格表可编辑", level=2)
    add_bullets(
        doc,
        [
            "在“价格表可编辑”中可直接修改已有价格。",
            "可编辑字段主要包括“收盘价”和“手工锁定”。",
            "点击“保存价格表修改”后，系统会以手工价格来源保存。",
            "被“手工锁定”的价格，AK 自动导入默认不会覆盖。"],
    )
    add_heading(doc, "7.4 AKShare 自动导入（可选）", level=2)
    add_bullets(
        doc,
        [
            "价格页提供“快速更新今收”和“AK自动导入”两类外部价格能力。",
            "若页面提示“未检测到 AKShare”，先在运行环境安装 `pip install akshare` 并刷新页面。",
            "“快速更新今收”只准备当天一条价格，需要用户再次点击“确认写入今日收盘价”。",
            "自动导入使用的是 AKShare 的“收盘价”字段，不使用即时价或结算价。"],
    )
    add_heading(doc, "7.5 一键清空所有价格", level=2)
    add_bullets(
        doc,
        [
            "页面底部提供“一键清空所有价格”。",
            "必须先勾选“我确认清空”才能执行。",
            "该操作影响所有策略组、所有品种的价格数据，执行前务必备份。"],
    )

    add_heading(doc, "8. 现货头寸仓库管理", level=1)
    add_paragraph(doc, "该页面用于维护现货买入、现货平仓、明细回溯和汇总清零，适用于有现货套保场景的策略组。")
    add_heading(doc, "8.1 A. 现货头寸汇总", level=2)
    add_bullets(
        doc,
        [
            "展示当前策略组下各现货名称的买入总量、已对冲数量、可用数量、可用均价、现货已实现盈亏和对冲总盈亏。",
            "当某个现货名称可用数量为 0 时，用户可以按名称进行清零隐藏。"],
    )
    add_heading(doc, "8.2 B. 现货买入增仓录入", level=2)
    add_numbers(
        doc,
        [
            "录入买入日期、现货名称、数量、买入价、操作人和备注。",
            "点击确认按钮后，系统会生成现货买入头寸（lot）。",
            "新增的现货头寸会进入汇总区和明细流水区。"],
    )
    add_heading(doc, "8.3 C. 纯现货平仓", level=2)
    add_bullets(
        doc,
        [
            "适用于未关联结构、只对现货头寸本身做平仓。",
            "系统会按现货头寸计算可平数量，并记录纯现货平仓流水。"],
    )
    add_heading(doc, "8.4 D. 明细流水", level=2)
    add_bullets(
        doc,
        [
            "页面下方可展开查看“现货买入流水”和“平仓记录”。",
            "可用于对账、排查和回看历史操作。"],
    )

    add_heading(doc, "9. 期权头寸仓库管理", level=1)
    add_paragraph(doc, "这是日常交易和仓位管理使用频率最高的页面，主要用于查看在库结构头寸、执行各类平仓和回看平仓明细。")
    add_heading(doc, "9.1 仓库头寸总览", level=2)
    add_numbers(
        doc,
        [
            "选择策略组后，页面会自动根据价格和结构计算当前仓库。",
            "在“仓库统计截至日期”选择仓库统计日期；该日期会影响仓库头寸和下方默认平仓日期。",
            "可用“品种快速筛选”聚焦某一个或多个合约月份。",
            "上方“A. 在库总览”显示多头、空头、净头寸、对冲数量、对冲价差收益等指标。"],
    )
    add_heading(doc, "9.2 B. 结构在库表（每结构一行，可快速平仓）", level=2)
    add_bullets(
        doc,
        [
            "表格按结构展示“可平数量、在库均价、平仓方向、平仓数量、平仓价格”等字段。",
            "可按结构、方向、风险子、品种进行多维筛选。",
            "勾选需要处理的结构后，可进入“批量快速平仓”。",
            "如果某条结构平仓后剩余可平数量为 0，系统会自动清理空结构，使其不再出现在活动页中。"],
    )
    add_heading(doc, "9.3 TRS 快捷管理（新增 + 换月）", level=2)
    add_bullets(
        doc,
        [
            "“B1.1 快速新增TRS头寸”：用于录入新的 TRS 仓位。",
            "“B1.2 TRS换月管理”：用于对旧合约 TRS 执行换月；系统会自动生成平旧开新的记录。",
            "若旧合约在换月后头寸归零，系统会自动删除旧 TRS 结构。"],
    )
    add_heading(doc, "9.4 其他平仓功能", level=2)
    add_bullets(
        doc,
        [
            "B2. 多空对称平仓（按结构配对）：适用于多头与空头结构按对称数量同时处理。",
            "B3. 现货对冲平仓（按现货头寸配对）：适用于结构平仓同时匹配现货头寸。",
            "C. 平仓录入（结构内 + 外部）：包括结构内头寸平仓、结构整体平仓、外部平仓录入。",
            "D. 平仓明细与回溯：可筛选查看已保存的平仓批次，并支持撤回批次。"],
    )
    add_heading(doc, "9.5 平仓与撤回的使用建议", level=2)
    add_numbers(
        doc,
        [
            "先确认仓库统计截至日期与业务日期一致，再执行平仓。",
            "批量快速平仓时，重点核对“统一平仓价格”“统一平仓数量”和行内明细。",
            "若平仓方向与结构方向不匹配、或数量超过当前可平数量，系统会阻止保存。",
            "若本次平仓保存成功但发现业务输入错误，可在“D. 平仓明细与回溯”中撤回对应批次。"],
    )

    add_heading(doc, "10. 监控计算", level=1)
    add_paragraph(doc, "监控计算页是统一结果页，用于查看价格完整性、日度口径、风险敞口、TRS 监控、平仓明细，并导出汇报图片。")
    add_heading(doc, "10.1 顶部全局参数", level=2)
    add_bullets(
        doc,
        [
            "顶部依次选择“策略组”“监控日期”“品种”，所有下方表格都受这三个条件联动。",
            "点击“重新计算并生成图片”可刷新监控结果和汇报图。",
            "页面上方会生成“汇报图片预览”，可点击“下载汇报图片 PNG”导出。"],
    )
    add_heading(doc, "10.2 价格完整性监控", level=2)
    add_bullets(
        doc,
        [
            "用于检查监控日期之前的价格是否完整。",
            "若缺少交易日价格，系统会列出结构、缺失日期列表和剩余观察信息。",
            "建议先修复价格缺口，再解读后续风险结果。"],
    )
    add_heading(doc, "10.3 日度维度监控软件", level=2)
    add_paragraph(doc, "下方 tab 按不同口径展示同一监控日结果。")
    add_table(
        doc,
        ["Tab 名称", "说明"],
        [
            ["结构监控总览", "展示各结构状态、剩余交易日、当前浮盈亏、雪球折价转期货等关键指标。"],
            ["结构日度明细", "展示非 TRS 结构的日度明细；TRS 已单独拆出，不在此 tab 中展示。"],
            ["策略组日度汇总", "按策略组/品种汇总当日生成量、净持仓、浮动盈亏、平仓盈亏和总盈亏。"],
            ["风险敞口区间", "查看结构和策略组层面的剩余最小/最大规模、敞口上下界等。"],
            ["TRS监控", "只展示仍有剩余头寸的 TRS；默认字段为结构详情、风险子、方向、入场价、数量、当日浮盈亏。"],
            ["平仓明细", "展示当前筛选下的结构平仓、现货对冲平仓等回溯明细，并支持导出。"],
        ],
    )
    add_heading(doc, "10.4 使用监控页的建议", level=2)
    add_bullets(
        doc,
        [
            "先看“价格完整性监控”，确认价格无缺口。",
            "再看“结构监控总览”和“风险敞口区间”，定位风险和剩余暴露。",
            "如需看 TRS，请直接切到“TRS监控”，不要再去“结构日度明细”寻找 TRS。",
            "如需对账平仓结果，使用“平仓明细”并导出 CSV。"],
    )

    add_heading(doc, "11. 常用业务场景操作指引", level=1)
    add_heading(doc, "11.1 新建一个普通累计结构", level=2)
    add_numbers(
        doc,
        [
            "进入“生成策略组”，确认目标策略组已存在。",
            "进入“结构录入”，选择策略组。",
            "填写结构编号、风险子、标的名称/代码、期权结构、报价日期、方向、起止日期、入场价等参数。",
            "点击“保存结构”。",
            "进入“价格录入”，按该结构区间补齐价格。",
            "进入“监控计算”，选择该策略组和监控日期，查看结构监控结果。"],
    )
    add_heading(doc, "11.2 新增一个 TRS 头寸并在监控页查看", level=2)
    add_numbers(
        doc,
        [
            "进入“期权头寸仓库管理”。",
            "在“TRS快捷管理（新增 + 换月）”中打开“B1.1 快速新增TRS头寸”。",
            "填写 TRS 的方向、品种、数量和入场价，保存新增记录。",
            "切换到“监控计算”，选择对应策略组和日期。",
            "在“TRS监控” tab 查看该 TRS 的结构详情、数量和当日浮盈亏。"],
    )
    add_heading(doc, "11.3 批量快速平仓并核对结果", level=2)
    add_numbers(
        doc,
        [
            "进入“期权头寸仓库管理”，在“B. 结构在库表”筛选目标结构。",
            "勾选要平仓的结构，点击“批量快速平仓”。",
            "确认平仓日期，录入统一平仓价和统一平仓数量，必要时逐行调整。",
            "点击保存并在确认弹窗中再次核对结构、方向、平仓数量和平仓价格。",
            "保存成功后，进入“D. 平仓明细与回溯”或“监控计算 -> 平仓明细”查看结果。"],
    )
    add_heading(doc, "11.4 生成日报图片", level=2)
    add_numbers(
        doc,
        [
            "进入“监控计算”。",
            "选择策略组、监控日期、品种。",
            "确认上方的价格完整性监控无明显问题。",
            "查看页面顶部生成的“汇报图片预览”。",
            "点击“下载汇报图片 PNG”导出。"],
    )

    add_heading(doc, "12. 常见问题与排查", level=1)
    add_table(
        doc,
        ["现象", "可能原因", "处理建议"],
        [
            ["监控页提示暂无结果", "尚未录入结构、价格，或监控日期之前没有有效计算数据", "先检查结构录入和价格录入，再重新进入监控页"],
            ["保存结构失败", "结构编号为空、与其他策略组重复，或保存后会立即触发终止", "修改结构编号；阅读提示信息并确认终止风险"],
            ["平仓保存失败", "平仓方向错误、数量超过可平数量、校验发现会加重超平仓", "回到仓库表核对方向、数量和日期"],
            ["AK 自动导入不可用", "未安装 AKShare", "安装 `pip install akshare` 后刷新页面"],
            ["监控页里找不到 TRS", "TRS 不在结构日度明细展示；只有有剩余头寸的 TRS 才出现在 TRS监控", "切换到“TRS监控”，并确认该 TRS 仍有头寸"],
            ["某条结构突然不见了", "结构已终止，或在平仓后剩余头寸归零被系统自动清理", "在终止结构列表或平仓明细中回看历史记录"],
        ],
    )

    add_heading(doc, "13. 数据安全与操作建议", level=1)
    add_bullets(
        doc,
        [
            "执行“覆盖导入”“一键清空所有价格”“删除策略组”“彻底删除结构”等动作前，先做导出备份。",
            "录入价格后，建议在价格表中对关键手工价格勾选“手工锁定”，避免自动导入覆盖。",
            "监控口径依赖价格完整性；一旦价格缺口较多，先修价再看风险。",
            "对外汇报优先使用“监控计算”页生成的汇报图片和 CSV，避免手工截图造成口径偏差。",
            "如果需要追溯最近一笔平仓，优先到“期权头寸仓库管理 -> D. 平仓明细与回溯”查看或撤回。"],
    )

    add_heading(doc, "14. 附录：按钮和页面速查", level=1)
    add_table(
        doc,
        ["页面", "常用按钮 / 区域", "用途"],
        [
            ["生成策略组", "创建/更新策略组", "保存策略组主数据"],
            ["生成策略组", "导出选中策略组 / 确认导入策略组", "导出或导入 JSON 数据包"],
            ["结构录入", "保存结构 / 下载结构报价图片 PNG", "保存结构并输出报价图"],
            ["结构录入", "保存表格修改 / 删除选中结构 / 确认彻底删除", "维护当前组内结构和已终止结构"],
            ["价格录入", "快速更新今收 / AK自动导入 / 保存结构区间价格", "维护批量价格和当日价格"],
            ["价格录入", "保存单日价格 / 保存价格表修改", "快速录入和表格维护价格"],
            ["现货头寸仓库管理", "现货买入增仓 / 纯现货平仓 / 明细流水", "维护现货库存和流水"],
            ["期权头寸仓库管理", "批量快速平仓 / 保存新增TRS头寸 / 保存TRS换月", "处理在库结构和平仓操作"],
            ["期权头寸仓库管理", "结构内头寸平仓 / 外部平仓录入 / 撤回批次", "执行手工平仓和回溯"],
            ["监控计算", "重新计算并生成图片 / 下载汇报图片 PNG", "刷新监控并导出日报图"],
        ],
    )

    add_paragraph(doc, "手册到此结束。建议将本文件与系统导出的策略组 JSON 一并保存，作为业务交接和培训材料。")


def main() -> None:
    doc = build_document()
    write_manual(doc)
    doc.save(OUT_PATH)
    doc.save(ALT_OUT_PATH)
    print(str(OUT_PATH))
    print(str(ALT_OUT_PATH))


if __name__ == "__main__":
    main()
